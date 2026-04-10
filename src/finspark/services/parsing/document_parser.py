"""Document parsing service - extracts structured data from BRDs, SOWs, API specs."""

import json
import re
from pathlib import Path
from typing import Any

import yaml

from finspark.schemas.common import DocType
from finspark.schemas.documents import (
    ExtractedAuth,
    ExtractedEndpoint,
    ExtractedField,
    ParsedDocumentResult,
)


class DocumentParser:
    """Parses various document formats and extracts integration requirements."""

    # Regex patterns for entity extraction
    URL_PATTERN = re.compile(r"https?://[^\s<>\"']+")
    API_PATH_PATTERN = re.compile(r"(?:GET|POST|PUT|DELETE|PATCH)\s+(/[a-zA-Z0-9/_\-{}]+)")
    ENDPOINT_PATTERN = re.compile(r"/(?:api|v\d+)/[a-zA-Z0-9/_\-{}]+")
    FIELD_PATTERN = re.compile(
        r"\b((?:applicant|customer|borrower|account|loan|pan|aadhaar|gstin|"
        r"mobile|email|address|name|dob|amount|score|status|reference|"
        r"transaction|payment|ifsc|vpa|upi)[_a-zA-Z]*)\b",
        re.IGNORECASE,
    )
    AUTH_KEYWORDS = re.compile(
        r"\b(api[_\s]?key|oauth|bearer|certificate|mTLS|basic\s*auth|jwt|token|secret)\b",
        re.IGNORECASE,
    )

    @staticmethod
    def _normalize_doc_type(doc_type: str) -> str:
        """Normalize doc_type to a valid DocType enum value, defaulting to 'brd'."""
        valid = {e.value for e in DocType}
        if doc_type not in valid:
            return "brd"
        return doc_type

    def parse(self, file_path: Path, doc_type: str = "auto") -> ParsedDocumentResult:
        """Parse a document and extract structured information."""
        doc_type = self._normalize_doc_type(doc_type)
        suffix = file_path.suffix.lower()

        if suffix == ".docx":
            return self._parse_docx(file_path, doc_type=doc_type)
        elif suffix == ".pdf":
            return self._parse_pdf(file_path, doc_type=doc_type)
        elif suffix in (".yaml", ".yml"):
            return self._parse_openapi(file_path)
        elif suffix == ".json":
            return self._parse_json_spec(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    def parse_text(self, text: str, doc_type: str = "brd") -> ParsedDocumentResult:
        """Parse raw text content and extract structured information."""
        doc_type = self._normalize_doc_type(doc_type)
        endpoints = self._extract_endpoints(text)
        fields = self._extract_fields(text)
        auth = self._extract_auth_requirements(text)
        services = self._extract_services(text)
        sections = self._extract_sections(text)
        security_reqs = self._extract_security_requirements(text)
        sla_reqs = self._extract_sla_requirements(text)

        total_entities = len(endpoints) + len(fields) + len(auth) + len(services)
        confidence = min(1.0, total_entities / 20.0)  # Normalize to 0-1

        return ParsedDocumentResult(
            doc_type=doc_type,
            title=self._extract_title(text),
            summary=self._extract_summary(text),
            services_identified=services,
            endpoints=endpoints,
            fields=fields,
            auth_requirements=auth,
            security_requirements=security_reqs,
            sla_requirements=sla_reqs,
            sections=sections,
            confidence_score=round(confidence, 2),
            raw_entities=self._extract_all_entities(text),
        )

    def build_result_from_llm(
        self, llm_data: dict[str, Any], doc_type: str, original_text: str
    ) -> ParsedDocumentResult:
        """Build ParsedDocumentResult from LLM extraction output.

        Augments LLM results with regex-extracted fields the LLM may have missed.
        """
        endpoints = [
            ExtractedEndpoint(
                path=ep.get("path", ""),
                method=ep.get("method", "GET"),
                description=ep.get("description", ""),
                parameters=[],
                is_mandatory=ep.get("is_mandatory", True),
            )
            for ep in llm_data.get("endpoints", [])
        ]

        fields = [
            ExtractedField(
                name=f.get("name", ""),
                data_type=f.get("data_type", "string"),
                is_required=f.get("is_required", False),
                source_section=f.get("source_section", ""),
            )
            for f in llm_data.get("fields", [])
        ]

        auth = [
            ExtractedAuth(
                auth_type=a.get("auth_type", "api_key"),
                details=a.get("details", {}),
            )
            for a in llm_data.get("auth_requirements", [])
        ]

        # Augment with regex-extracted fields the LLM may have missed
        regex_fields = self._extract_fields(original_text)
        llm_field_names = {f.name for f in fields}
        for rf in regex_fields:
            if rf.name not in llm_field_names:
                fields.append(rf)

        total_entities = len(endpoints) + len(fields) + len(auth) + len(
            llm_data.get("services_identified", [])
        )
        confidence = min(1.0, max(0.7, total_entities / 20.0))

        sla_raw = llm_data.get("sla_requirements", {})
        sla_list: list[str] = []
        if isinstance(sla_raw, dict):
            if sla_raw.get("response_time_ms"):
                sla_list.append(f"Response time: {sla_raw['response_time_ms']}ms")
            if sla_raw.get("availability_percent"):
                sla_list.append(f"Availability: {sla_raw['availability_percent']}%")

        return ParsedDocumentResult(
            doc_type=doc_type,
            title=llm_data.get("title", self._extract_title(original_text)),
            summary=llm_data.get("summary", self._extract_summary(original_text)),
            services_identified=llm_data.get("services_identified", []),
            endpoints=endpoints,
            fields=fields,
            auth_requirements=auth,
            security_requirements=llm_data.get("security_requirements", []),
            sla_requirements=sla_list,
            sections=self._extract_sections(original_text),
            confidence_score=round(confidence, 2),
            raw_entities=self._extract_all_entities(original_text),
        )

    def _parse_docx(self, file_path: Path, doc_type: str = "brd") -> ParsedDocumentResult:
        from docx import Document

        doc = Document(str(file_path))
        full_text_parts: list[str] = []
        for para in doc.paragraphs:
            full_text_parts.append(para.text)

        # Extract table data
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                full_text_parts.append(row_text)

        full_text = "\n".join(full_text_parts)
        return self.parse_text(full_text, doc_type=doc_type)

    def _parse_pdf(self, file_path: Path, doc_type: str = "brd") -> ParsedDocumentResult:
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        text_parts: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        full_text = "\n".join(text_parts)
        return self.parse_text(full_text, doc_type=doc_type)

    def _parse_openapi(self, file_path: Path) -> ParsedDocumentResult:
        with open(file_path) as f:
            spec = yaml.safe_load(f)
        return self._parse_openapi_dict(spec)

    def _parse_json_spec(self, file_path: Path) -> ParsedDocumentResult:
        with open(file_path) as f:
            data = json.load(f)

        # Check if it's an OpenAPI spec
        if "openapi" in data or "swagger" in data:
            return self._parse_openapi_dict(data)

        # Treat as a generic JSON config
        return ParsedDocumentResult(
            doc_type="api_spec",
            title=data.get("title", "JSON Specification"),
            summary=data.get("description", ""),
            confidence_score=0.8,
        )

    @staticmethod
    def _resolve_ref(ref: str, spec: dict[str, Any]) -> dict[str, Any]:
        """Resolve a JSON $ref string against the spec (local refs only)."""
        # Expected format: "#/components/schemas/FooBar"
        if not ref.startswith("#/"):
            return {}
        parts = ref.removeprefix("#/").split("/")
        node: Any = spec
        for part in parts:
            if not isinstance(node, dict):
                return {}
            node = node.get(part, {})
        return node if isinstance(node, dict) else {}

    def _resolve_schema(self, schema: dict[str, Any], spec: dict[str, Any]) -> dict[str, Any]:
        """Return the schema dict with $ref resolved, if present."""
        if "$ref" in schema:
            return self._resolve_ref(schema["$ref"], spec)
        return schema

    def _parse_openapi_dict(self, spec: dict[str, Any]) -> ParsedDocumentResult:
        """Parse an OpenAPI specification dictionary."""
        endpoints: list[ExtractedEndpoint] = []
        fields: list[ExtractedField] = []

        info = spec.get("info", {})
        paths = spec.get("paths", {})
        servers = spec.get("servers", [])

        # Extract endpoints and inline request/response fields
        for path, methods in paths.items():
            for method, details in methods.items():
                if method not in ("get", "post", "put", "delete", "patch"):
                    continue

                params = []
                for p in details.get("parameters", []):
                    params.append(
                        {
                            "name": p.get("name", ""),
                            "in": p.get("in", "query"),
                            "required": str(p.get("required", False)),
                        }
                    )

                endpoints.append(
                    ExtractedEndpoint(
                        path=path,
                        method=method.upper(),
                        description=details.get("summary", details.get("description", "")),
                        parameters=params,
                        is_mandatory=True,
                    )
                )

                # Extract fields from inline requestBody schema
                req_body = details.get("requestBody", {})
                req_schema = self._resolve_schema(
                    req_body.get("content", {})
                    .get("application/json", {})
                    .get("schema", {}),
                    spec,
                )
                if req_schema.get("properties"):
                    section = f"{method.upper()} {path} request"
                    req_required = set(req_schema.get("required", []))
                    for fname, fdef in req_schema["properties"].items():
                        fdef = self._resolve_schema(fdef, spec)
                        fields.append(
                            ExtractedField(
                                name=fname,
                                data_type=fdef.get("type", "string"),
                                description=fdef.get("description", ""),
                                is_required=fname in req_required,
                                sample_value=str(fdef.get("example", "")),
                                source_section=section,
                            )
                        )

                # Extract fields from inline response schema (200/201/202)
                for status_code in ("200", "201", "202", 200, 201, 202):
                    resp_schema = self._resolve_schema(
                        details.get("responses", {})
                        .get(status_code, {})
                        .get("content", {})
                        .get("application/json", {})
                        .get("schema", {}),
                        spec,
                    )
                    if resp_schema.get("properties"):
                        section = f"{method.upper()} {path} response"
                        for fname, fdef in resp_schema["properties"].items():
                            fdef = self._resolve_schema(fdef, spec)
                            fields.append(
                                ExtractedField(
                                    name=fname,
                                    data_type=fdef.get("type", "string"),
                                    description=fdef.get("description", ""),
                                    is_required=False,
                                    sample_value=str(fdef.get("example", "")),
                                    source_section=section,
                                )
                            )

        # Extract fields from components.schemas (named schemas)
        schemas = spec.get("components", {}).get("schemas", {})
        for schema_name, schema_def in schemas.items():
            required_fields = set(schema_def.get("required", []))
            properties = schema_def.get("properties", {})
            for field_name, field_def in properties.items():
                fields.append(
                    ExtractedField(
                        name=f"{schema_name}.{field_name}",
                        data_type=field_def.get("type", "string"),
                        description=field_def.get("description", ""),
                        is_required=field_name in required_fields,
                        sample_value=str(field_def.get("example", "")),
                        source_section=schema_name,
                    )
                )

        # Deduplicate fields by name (keep first occurrence)
        seen_fields: set[str] = set()
        unique_fields: list[ExtractedField] = []
        for f in fields:
            if f.name not in seen_fields:
                seen_fields.add(f.name)
                unique_fields.append(f)
        fields = unique_fields

        # Extract auth
        auth_reqs: list[ExtractedAuth] = []
        security_schemes = spec.get("components", {}).get("securitySchemes", {})
        for scheme_name, scheme_def in security_schemes.items():
            auth_reqs.append(
                ExtractedAuth(
                    auth_type=scheme_def.get("type", "apiKey"),
                    details={
                        "name": scheme_name,
                        "scheme": scheme_def.get("scheme", ""),
                        "in": scheme_def.get("in", ""),
                    },
                )
            )

        base_urls = [s.get("url", "") for s in servers]
        services = [info.get("title", "API Service")]

        return ParsedDocumentResult(
            doc_type="api_spec",
            title=info.get("title", ""),
            summary=info.get("description", ""),
            services_identified=services,
            endpoints=endpoints,
            fields=fields,
            auth_requirements=auth_reqs,
            sections={"base_urls": ", ".join(base_urls)},
            confidence_score=0.95,
        )

    def _extract_endpoints(self, text: str) -> list[ExtractedEndpoint]:
        endpoints = []
        seen = set()

        # Match explicit HTTP method + path
        for match in self.API_PATH_PATTERN.finditer(text):
            path = match.group(1)
            method = match.group(0).split()[0]
            if path not in seen:
                seen.add(path)
                endpoints.append(ExtractedEndpoint(path=path, method=method))

        # Match standalone API paths
        for match in self.ENDPOINT_PATTERN.finditer(text):
            path = match.group(0)
            if path not in seen:
                seen.add(path)
                endpoints.append(ExtractedEndpoint(path=path, method="GET"))

        return endpoints

    def _extract_fields(self, text: str) -> list[ExtractedField]:
        fields = []
        seen: set[str] = set()

        for match in self.FIELD_PATTERN.finditer(text):
            name = match.group(1).lower()
            if name not in seen and len(name) > 3:
                seen.add(name)
                # Infer type from name
                data_type = self._infer_field_type(name)
                fields.append(
                    ExtractedField(
                        name=name,
                        data_type=data_type,
                        is_required=True,
                    )
                )

        return fields

    def _extract_auth_requirements(self, text: str) -> list[ExtractedAuth]:
        auth_reqs = []
        seen: set[str] = set()

        for match in self.AUTH_KEYWORDS.finditer(text):
            auth_type = match.group(1).lower().replace(" ", "_")
            if auth_type not in seen:
                seen.add(auth_type)
                auth_reqs.append(ExtractedAuth(auth_type=auth_type))

        return auth_reqs

    def _extract_services(self, text: str) -> list[str]:
        service_keywords = [
            "CIBIL",
            "Experian",
            "CRIF",
            "Equifax",
            "Aadhaar",
            "PAN",
            "DigiLocker",
            "eKYC",
            "KYC",
            "GST",
            "GSTN",
            "Razorpay",
            "PayU",
            "NPCI",
            "UPI",
            "NEFT",
            "IMPS",
            "RTGS",
            "Account Aggregator",
            "fraud detection",
            "fraud engine",
            "SMS gateway",
            "email gateway",
        ]
        found = []
        text_upper = text.upper()
        for keyword in service_keywords:
            if keyword.upper() in text_upper:
                found.append(keyword)
        return list(set(found))

    def _extract_sections(self, text: str) -> dict[str, str]:
        """Extract document sections based on common headers."""
        sections: dict[str, str] = {}
        current_section = "overview"
        current_content: list[str] = []

        section_patterns = [
            (r"(?i)(?:project|executive)\s*(?:overview|summary)", "overview"),
            (r"(?i)integration\s*requirements?", "integration_requirements"),
            (r"(?i)(?:data\s*flow|architecture)", "data_flow"),
            (r"(?i)security\s*requirements?", "security"),
            (r"(?i)(?:sla|performance)\s*requirements?", "sla"),
            (r"(?i)(?:error|exception)\s*handling", "error_handling"),
            (r"(?i)test(?:ing)?\s*requirements?", "testing"),
            (r"(?i)field\s*(?:mapping|definition)", "field_mapping"),
        ]

        for line in text.split("\n"):
            matched = False
            for pattern, section_name in section_patterns:
                if re.search(pattern, line):
                    if current_content:
                        sections[current_section] = "\n".join(current_content).strip()
                    current_section = section_name
                    current_content = []
                    matched = True
                    break
            if not matched:
                current_content.append(line)

        if current_content:
            sections[current_section] = "\n".join(current_content).strip()

        return sections

    def _extract_security_requirements(self, text: str) -> list[str]:
        security_keywords = [
            r"encrypt(?:ion|ed)",
            r"PII\s*(?:mask|protect|handl)",
            r"audit\s*(?:log|trail)",
            r"(?:data|information)\s*security",
            r"(?:PCI|SOX|GDPR|RBI)\s*(?:compliance|DSS)",
            r"access\s*control",
            r"data\s*(?:masking|anonymi)",
        ]
        found = []
        for pattern in security_keywords:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                found.append(match.group(0))
        return found

    def _extract_sla_requirements(self, text: str) -> dict[str, str]:
        sla: dict[str, str] = {}
        # Look for response time patterns
        time_match = re.search(
            r"(?:response\s*time|latency)[:\s]*(\d+\s*(?:ms|seconds?))", text, re.IGNORECASE
        )
        if time_match:
            sla["response_time"] = time_match.group(1)

        avail_match = re.search(r"(?:availability|uptime)[:\s]*([\d.]+%)", text, re.IGNORECASE)
        if avail_match:
            sla["availability"] = avail_match.group(1)

        return sla

    def _extract_title(self, text: str) -> str:
        _SEPARATOR_RE = re.compile(r"^[|\-#=~\s]+$")
        lines = text.strip().split("\n")
        for line in lines[:5]:
            line = line.strip()
            if not line:
                continue
            # Skip table separators and markdown heading/rule characters
            if line[0] in ("|", "-", "#", "=", "~"):
                continue
            if _SEPARATOR_RE.match(line):
                continue
            if len(line) > 10 and len(line) < 200:
                return line
        return "Untitled Document"

    def _extract_summary(self, text: str) -> str:
        sentences = re.split(r"[.!?]\s+", text[:2000])
        if len(sentences) >= 3:
            return ". ".join(sentences[:3]) + "."
        return text[:500]

    def _extract_all_entities(self, text: str) -> list[str]:
        entities: set[str] = set()
        for match in self.URL_PATTERN.finditer(text):
            entities.add(match.group(0))
        for match in self.ENDPOINT_PATTERN.finditer(text):
            entities.add(match.group(0))
        for match in self.FIELD_PATTERN.finditer(text):
            entities.add(match.group(1))
        return sorted(entities)[:50]

    @staticmethod
    def _infer_field_type(name: str) -> str:
        name_lower = name.lower()
        if any(k in name_lower for k in ("date", "dob", "created", "updated")):
            return "date"
        if any(k in name_lower for k in ("amount", "score", "balance", "rate")):
            return "number"
        if any(k in name_lower for k in ("email",)):
            return "email"
        if any(k in name_lower for k in ("mobile", "phone")):
            return "phone"
        if any(k in name_lower for k in ("is_", "has_", "active", "enabled")):
            return "boolean"
        return "string"
