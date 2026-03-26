#!/usr/bin/env python3
"""Seed additional test data into FinSpark.

Creates sample configurations for multiple adapters, runs simulations,
and generates audit trail entries for a realistic demo environment.

Usage:
    python scripts/seed_data.py [--base-url http://localhost:8000]
"""

from __future__ import annotations

import argparse
import json
import sys
import tempfile
import time
from pathlib import Path
from typing import Any

import httpx
from docx import Document

# ── Colors ────────────────────────────────────────────────────────────────────
RESET = "\033[0m"
BOLD = "\033[1m"
DIM = "\033[2m"
GREEN = "\033[92m"
YELLOW = "\033[93m"
CYAN = "\033[96m"
RED = "\033[91m"


def ok(msg: str) -> None:
    print(f"  {GREEN}✓{RESET} {msg}")


def info(msg: str) -> None:
    print(f"  {CYAN}ℹ{RESET} {msg}")


def fail(msg: str) -> None:
    print(f"  {RED}✗{RESET} {msg}")


def section(title: str) -> None:
    print(f"\n{CYAN}{BOLD}── {title} ──{RESET}")


# ── BRD generators for different adapters ─────────────────────────────────────

BRD_SPECS: list[dict[str, Any]] = [
    {
        "name": "eKYC Aadhaar Verification",
        "adapter_keyword": "ekyc",
        "doc_type": "brd",
        "filename": "eKYC_Integration_BRD.docx",
        "config_name": "eKYC-Aadhaar-Integration",
        "sections": {
            "overview": (
                "This BRD covers the integration of Aadhaar-based electronic KYC (eKYC) "
                "verification services. The system must support OTP-based and biometric-based "
                "Aadhaar verification, PAN card validation, and DigiLocker document retrieval."
            ),
            "endpoints": [
                ("POST", "/verify/aadhaar", "Verify Aadhaar number via OTP"),
                ("POST", "/verify/pan", "Verify PAN card details"),
                ("POST", "/digilocker/fetch", "Fetch documents from DigiLocker"),
            ],
            "fields": [
                ("aadhaar_number", "string", True),
                ("customer_name", "string", True),
                ("pan_number", "string", False),
                ("date_of_birth", "date", True),
                ("mobile_number", "string", True),
                ("consent", "boolean", True),
            ],
            "auth_type": "api_key",
            "security": [
                "Aadhaar numbers must be encrypted and never logged in plaintext",
                "Biometric data must not be stored beyond the verification session",
                "All eKYC requests must carry valid consent tokens",
            ],
        },
    },
    {
        "name": "GST Verification",
        "adapter_keyword": "gst",
        "doc_type": "brd",
        "filename": "GST_Verification_BRD.docx",
        "config_name": "GST-Verification-Integration",
        "sections": {
            "overview": (
                "Integration requirements for GST verification services including GSTIN "
                "validation, return filing status checks, and taxpayer profile retrieval. "
                "Required for MSME lending workflows to verify business legitimacy."
            ),
            "endpoints": [
                ("POST", "/verify/gstin", "Validate GSTIN and fetch details"),
                ("GET", "/returns/status", "Check GST return filing compliance"),
                ("GET", "/profile", "Fetch taxpayer business profile"),
            ],
            "fields": [
                ("gstin", "string", True),
                ("financial_year", "string", False),
                ("return_type", "string", False),
                ("business_pan", "string", True),
            ],
            "auth_type": "api_key",
            "security": [
                "GST data is business-sensitive and must be access-controlled",
                "API rate limits: 50 requests/minute per tenant",
            ],
        },
    },
    {
        "name": "Payment Gateway Disbursement",
        "adapter_keyword": "payment",
        "doc_type": "brd",
        "filename": "Payment_Gateway_BRD.docx",
        "config_name": "Payment-Disbursement-Integration",
        "sections": {
            "overview": (
                "This BRD covers loan disbursement via the payment gateway. The integration "
                "supports NEFT, IMPS, RTGS, and UPI payment modes. It handles payment creation, "
                "status tracking, bank transfers, and refund processing."
            ),
            "endpoints": [
                ("POST", "/payments/create", "Initiate a loan disbursement payment"),
                ("GET", "/payments/{id}", "Track payment status"),
                ("POST", "/transfers/create", "Create bank-to-bank transfer"),
                ("POST", "/refunds/create", "Process loan refund"),
            ],
            "fields": [
                ("amount", "number", True),
                ("account_number", "string", True),
                ("ifsc_code", "string", True),
                ("beneficiary_name", "string", True),
                ("reference_id", "string", True),
                ("payment_mode", "string", True),
                ("vpa", "string", False),
            ],
            "auth_type": "api_key",
            "security": [
                "Account numbers must be masked in logs (show last 4 digits only)",
                "Payment amounts above 10L INR require dual-authorization",
                "All disbursements must be reconciled within T+1",
            ],
        },
    },
    {
        "name": "Fraud Detection Engine",
        "adapter_keyword": "fraud",
        "doc_type": "brd",
        "filename": "Fraud_Detection_BRD.docx",
        "config_name": "Fraud-Detection-Integration",
        "sections": {
            "overview": (
                "Integration with real-time fraud detection engine for loan application "
                "screening. The system must perform device fingerprinting, velocity checks, "
                "and fraud risk scoring before loan approval. Scores above threshold 0.7 "
                "trigger manual review."
            ),
            "endpoints": [
                ("POST", "/score", "Get fraud risk score for application"),
                ("POST", "/verify/device", "Device fingerprint analysis"),
                ("POST", "/verify/velocity", "Check transaction velocity patterns"),
            ],
            "fields": [
                ("customer_id", "string", True),
                ("transaction_amount", "number", True),
                ("device_id", "string", False),
                ("ip_address", "string", False),
                ("mobile_number", "string", False),
                ("email_address", "string", False),
            ],
            "auth_type": "api_key",
            "security": [
                "IP addresses and device fingerprints are PII — encrypt at rest",
                "Fraud scores must not be exposed to end customers",
                "Maintain 90-day rolling window of fraud signals for model training",
            ],
        },
    },
    {
        "name": "SMS Notification Gateway",
        "adapter_keyword": "sms",
        "doc_type": "brd",
        "filename": "SMS_Gateway_BRD.docx",
        "config_name": "SMS-Notification-Integration",
        "sections": {
            "overview": (
                "Integration with SMS gateway for transactional notifications throughout "
                "the loan lifecycle — application received, approved, disbursed, EMI due, "
                "EMI overdue. Must support DLT-registered templates as per TRAI regulations."
            ),
            "endpoints": [
                ("POST", "/send", "Send transactional SMS"),
                ("GET", "/status/{id}", "Check SMS delivery status"),
                ("GET", "/templates", "List registered DLT templates"),
            ],
            "fields": [
                ("mobile_number", "string", True),
                ("message", "string", True),
                ("template_id", "string", True),
                ("sender_id", "string", True),
            ],
            "auth_type": "api_key",
            "security": [
                "SMS content must not contain full account numbers or PAN",
                "Delivery reports must be retained for 30 days",
                "Rate limit: 100 SMS/second per sender ID",
            ],
        },
    },
]


def create_simple_brd(spec: dict[str, Any], output_path: Path) -> Path:
    """Create a simple BRD docx for the given spec."""
    doc = Document()
    sections = spec["sections"]

    doc.add_heading(f"Business Requirements Document: {spec['name']}", level=0)

    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(sections["overview"])

    doc.add_heading("2. API Endpoints", level=1)
    for method, path, desc in sections["endpoints"]:
        doc.add_paragraph(f"{method} {path} — {desc}", style="List Bullet")

    doc.add_heading("3. Field Requirements", level=1)
    for name, dtype, required in sections["fields"]:
        req_str = "Required" if required else "Optional"
        doc.add_paragraph(f"{name} ({dtype}) — {req_str}", style="List Bullet")

    doc.add_heading("4. Authentication", level=1)
    doc.add_paragraph(f"Authentication type: {sections['auth_type']}")

    doc.add_heading("5. Security Requirements", level=1)
    for sec in sections["security"]:
        doc.add_paragraph(sec, style="List Bullet")

    doc.save(str(output_path))
    return output_path


# ── API helpers ───────────────────────────────────────────────────────────────

class SeedClient:
    def __init__(self, base_url: str) -> None:
        self.api = f"{base_url.rstrip('/')}/api/v1"
        self.headers = {
            "X-Tenant-ID": "demo-tenant",
            "X-Tenant-Name": "Demo Bank",
            "X-Tenant-Role": "admin",
        }
        self.client = httpx.Client(headers=self.headers, timeout=30.0)

    def upload(self, path: Path, doc_type: str = "brd") -> dict[str, Any]:
        with open(path, "rb") as f:
            resp = self.client.post(
                f"{self.api}/documents/upload",
                params={"doc_type": doc_type},
                files={"file": (path.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
        resp.raise_for_status()
        return resp.json()

    def list_adapters(self) -> list[dict[str, Any]]:
        resp = self.client.get(f"{self.api}/adapters/")
        resp.raise_for_status()
        return resp.json()["data"]["adapters"]

    def generate_config(self, doc_id: str, av_id: str, name: str) -> dict[str, Any]:
        resp = self.client.post(
            f"{self.api}/configurations/generate",
            json={"document_id": doc_id, "adapter_version_id": av_id, "name": name},
        )
        resp.raise_for_status()
        return resp.json()

    def run_sim(self, config_id: str, test_type: str = "full") -> dict[str, Any]:
        resp = self.client.post(
            f"{self.api}/simulations/run",
            json={"configuration_id": config_id, "test_type": test_type},
        )
        resp.raise_for_status()
        return resp.json()

    def validate(self, config_id: str) -> dict[str, Any]:
        resp = self.client.post(f"{self.api}/configurations/{config_id}/validate")
        resp.raise_for_status()
        return resp.json()

    def close(self) -> None:
        self.client.close()


# ── Main seeding logic ────────────────────────────────────────────────────────

def seed(base_url: str) -> None:
    print(f"\n{CYAN}{BOLD}FinSpark Data Seeder{RESET}")
    print(f"{DIM}{'─' * 50}{RESET}")

    client = SeedClient(base_url)

    try:
        # Fetch adapters once
        adapters = client.list_adapters()
        adapter_map: dict[str, tuple[str, str]] = {}  # keyword -> (adapter_id, first_version_id)
        for a in adapters:
            keyword = a["name"].lower()
            versions = a.get("versions", [])
            if versions:
                adapter_map[keyword] = (a["id"], versions[0]["id"])

        info(f"Found {len(adapters)} adapters in registry")

        results: list[dict[str, str]] = []

        with tempfile.TemporaryDirectory() as tmpdir:
            for spec in BRD_SPECS:
                section(spec["name"])

                # Find matching adapter
                av_id: str | None = None
                adapter_id: str | None = None
                for key, (aid, vid) in adapter_map.items():
                    if spec["adapter_keyword"] in key:
                        adapter_id = aid
                        av_id = vid
                        break

                if not av_id:
                    fail(f"No adapter matched keyword '{spec['adapter_keyword']}' — skipping")
                    continue

                # Create BRD
                brd_path = Path(tmpdir) / spec["filename"]
                create_simple_brd(spec, brd_path)
                ok(f"Created BRD: {spec['filename']}")

                # Upload
                upload_resp = client.upload(brd_path, doc_type=spec["doc_type"])
                doc_id = upload_resp["data"]["id"]
                ok(f"Uploaded — doc_id={doc_id[:12]}...")

                # Generate config
                config_resp = client.generate_config(doc_id, av_id, spec["config_name"])
                config_id = config_resp["data"]["id"]
                mappings_count = len(config_resp["data"].get("field_mappings", []))
                ok(f"Config generated — {mappings_count} field mappings")

                # Validate
                val_resp = client.validate(config_id)
                coverage = val_resp["data"].get("coverage_score", 0)
                valid = val_resp["data"]["is_valid"]
                status_str = f"{GREEN}valid{RESET}" if valid else f"{YELLOW}issues{RESET}"
                ok(f"Validation: {status_str} (coverage {coverage:.0%})")

                # Run simulations (full + smoke)
                for test_type in ("full", "smoke"):
                    sim_resp = client.run_sim(config_id, test_type=test_type)
                    sim = sim_resp["data"]
                    sim_status = sim["status"]
                    color = GREEN if sim_status == "passed" else RED
                    ok(f"Simulation ({test_type}): {color}{sim_status}{RESET} — {sim['passed_tests']}/{sim['total_tests']}")

                results.append({
                    "adapter": spec["name"],
                    "config": spec["config_name"],
                    "config_id": config_id,
                    "doc_id": doc_id,
                })

        # Summary
        section("Seed Complete")
        print()
        for r in results:
            print(f"  {GREEN}●{RESET} {r['adapter']}")
            print(f"    Config: {DIM}{r['config_id']}{RESET}")
            print(f"    Doc:    {DIM}{r['doc_id']}{RESET}")

        total_configs = len(results)
        total_sims = total_configs * 2  # full + smoke per config
        print(f"\n  {BOLD}Totals:{RESET}")
        print(f"    Documents uploaded:   {total_configs}")
        print(f"    Configs generated:    {total_configs}")
        print(f"    Simulations run:      {total_sims}")
        print(f"    Audit entries:        ~{total_configs * 4 + total_sims}")
        print()

    except httpx.ConnectError:
        fail(f"Cannot connect to {base_url}")
        fail("Start the backend first: make dev")
        sys.exit(1)
    except httpx.HTTPStatusError as e:
        fail(f"HTTP {e.response.status_code}: {e.response.text[:300]}")
        sys.exit(1)
    finally:
        client.close()


def main() -> None:
    parser = argparse.ArgumentParser(description="Seed test data into FinSpark")
    parser.add_argument("--base-url", default="http://localhost:8000", help="Backend base URL")
    args = parser.parse_args()
    seed(args.base_url)


if __name__ == "__main__":
    main()
