#!/usr/bin/env python3
"""FinSpark Platform Demo - Full E2E walkthrough.

Demonstrates the complete integration workflow:
  BRD Upload -> Parsing -> Adapter Selection -> Config Generation -> Simulation

Usage:
    python scripts/demo.py [--base-url http://localhost:8000]
"""

from __future__ import annotations

import argparse
import json
import sys
import tempfile
import time
from pathlib import Path
from typing import Any

import httpx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ──────────────────────────────────────────────────────────────────────────────
# ANSI colors
# ──────────────────────────────────────────────────────────────────────────────
RESET = "\033[0m"
BOLD = "\033[1m"
DIM = "\033[2m"
RED = "\033[91m"
GREEN = "\033[92m"
YELLOW = "\033[93m"
BLUE = "\033[94m"
MAGENTA = "\033[95m"
CYAN = "\033[96m"
WHITE = "\033[97m"
BG_GREEN = "\033[42m"
BG_RED = "\033[41m"
BG_BLUE = "\033[44m"
BG_YELLOW = "\033[43m"


def banner() -> None:
    print(f"""
{CYAN}{BOLD}
    ╔═══════════════════════════════════════════════════════════════╗
    ║                                                               ║
    ║   ███████╗██╗███╗   ██╗███████╗██████╗  █████╗ ██████╗ ██╗  ║
    ║   ██╔════╝██║████╗  ██║██╔════╝██╔══██╗██╔══██╗██╔══██╗██║  ║
    ║   █████╗  ██║██╔██╗ ██║███████╗██████╔╝███████║██████╔╝██║  ║
    ║   ██╔══╝  ██║██║╚██╗██║╚════██║██╔═══╝ ██╔══██║██╔══██╗██║  ║
    ║   ██║     ██║██║ ╚████║███████║██║     ██║  ██║██║  ██║██║  ║
    ║   ╚═╝     ╚═╝╚═╝  ╚═══╝╚══════╝╚═╝     ╚═╝  ╚═╝╚═╝  ╚═╝╚═╝  ║
    ║                                                               ║
    ║   AI-Assisted Integration Configuration & Orchestration       ║
    ║   ─────────────────────────────────────────────────────────   ║
    ║   Full Platform Demo                                          ║
    ║                                                               ║
    ╚═══════════════════════════════════════════════════════════════╝
{RESET}""")


def step_header(step_num: int, title: str) -> None:
    print(f"\n{BG_BLUE}{WHITE}{BOLD} STEP {step_num} {RESET} {BLUE}{BOLD}{title}{RESET}")
    print(f"{DIM}{'─' * 70}{RESET}")


def success(msg: str) -> None:
    print(f"  {GREEN}✓{RESET} {msg}")


def info(msg: str) -> None:
    print(f"  {CYAN}ℹ{RESET} {msg}")


def warn(msg: str) -> None:
    print(f"  {YELLOW}⚠{RESET} {msg}")


def error(msg: str) -> None:
    print(f"  {RED}✗{RESET} {msg}")


def pretty_json(data: Any, indent: int = 4) -> str:
    return json.dumps(data, indent=indent, default=str)


def print_json_block(data: Any, max_lines: int = 25) -> None:
    formatted = pretty_json(data)
    lines = formatted.split("\n")
    for line in lines[:max_lines]:
        print(f"  {DIM}│{RESET} {line}")
    if len(lines) > max_lines:
        print(f"  {DIM}│ ... ({len(lines) - max_lines} more lines){RESET}")


def print_table(headers: list[str], rows: list[list[str]]) -> None:
    """Print a simple ASCII table."""
    col_widths = [len(h) for h in headers]
    for row in rows:
        for i, cell in enumerate(row):
            if i < len(col_widths):
                col_widths[i] = max(col_widths[i], len(str(cell)))

    header_line = " │ ".join(h.ljust(col_widths[i]) for i, h in enumerate(headers))
    separator = "─┼─".join("─" * col_widths[i] for i in range(len(headers)))

    print(f"  {DIM}┌─{'─┬─'.join('─' * col_widths[i] for i in range(len(headers)))}─┐{RESET}")
    print(f"  {DIM}│{RESET} {BOLD}{header_line}{RESET} {DIM}│{RESET}")
    print(f"  {DIM}├─{separator}─┤{RESET}")
    for row in rows:
        cells = " │ ".join(str(row[i]).ljust(col_widths[i]) if i < len(row) else " " * col_widths[i] for i in range(len(headers)))
        print(f"  {DIM}│{RESET} {cells} {DIM}│{RESET}")
    print(f"  {DIM}└─{'─┴─'.join('─' * col_widths[i] for i in range(len(headers)))}─┘{RESET}")


# ──────────────────────────────────────────────────────────────────────────────
# BRD Document Generator
# ──────────────────────────────────────────────────────────────────────────────

def create_brd_docx(output_path: Path) -> Path:
    """Create a realistic CIBIL Bureau Integration BRD document."""
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # -- Title Page --
    title = doc.add_heading("Business Requirements Document", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("CIBIL Credit Bureau Integration")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("\n\nProject: LendFlow Digital Lending Platform\n").font.size = Pt(12)
    meta.add_run("Version: 2.1\n").font.size = Pt(11)
    meta.add_run("Date: March 2026\n").font.size = Pt(11)
    meta.add_run("Classification: CONFIDENTIAL\n").font.size = Pt(11)
    meta.add_run("Author: Integration Engineering Team\n").font.size = Pt(11)

    doc.add_page_break()

    # -- 1. Project Overview --
    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "This document outlines the business requirements for integrating TransUnion CIBIL "
        "credit bureau services into the LendFlow Digital Lending Platform. The integration "
        "will enable real-time credit score retrieval, detailed credit report fetching, and "
        "bulk inquiry processing for loan underwriting workflows."
    )
    doc.add_paragraph(
        "The integration must support both individual and batch credit inquiries, handle "
        "consent verification per RBI guidelines, and maintain complete audit trails for "
        "regulatory compliance. The system processes approximately 50,000 credit inquiries "
        "per day across personal loans, home loans, and business lending products."
    )

    doc.add_heading("1.1 Business Objectives", level=2)
    objectives = [
        "Reduce credit assessment turnaround time from 48 hours to under 30 seconds",
        "Achieve 99.95% uptime for credit bureau connectivity",
        "Support concurrent processing of 500+ credit inquiries per minute",
        "Maintain full regulatory compliance with RBI and CIBIL guidelines",
        "Enable seamless version migration between CIBIL API v1 and v2",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style="List Bullet")

    doc.add_heading("1.2 Scope", level=2)
    doc.add_paragraph(
        "In-scope: Credit score retrieval, credit report fetching, bulk inquiry processing, "
        "consent management, error handling and retry logic, audit logging. "
        "Out-of-scope: Credit score model customization, direct CIBIL portal access, "
        "consumer dispute resolution workflows."
    )

    # -- 2. Integration Requirements --
    doc.add_heading("2. Integration Requirements", level=1)

    doc.add_heading("2.1 API Endpoints", level=2)
    doc.add_paragraph(
        "The following CIBIL API endpoints must be integrated into the lending platform:"
    )

    endpoint_table = doc.add_table(rows=5, cols=4)
    endpoint_table.style = "Light Grid Accent 1"
    hdr = endpoint_table.rows[0].cells
    hdr[0].text = "Endpoint"
    hdr[1].text = "Method"
    hdr[2].text = "Description"
    hdr[3].text = "Priority"

    endpoints = [
        ("/credit-score", "POST", "Fetch CIBIL score for individual applicant", "P0 - Critical"),
        ("/credit-report", "POST", "Fetch detailed credit report with trade lines", "P0 - Critical"),
        ("/bulk-inquiry", "POST", "Batch credit inquiry for multiple applicants", "P1 - High"),
        ("/consent/verify", "POST", "Verify borrower consent before inquiry", "P0 - Critical"),
    ]
    for i, (path, method, desc, priority) in enumerate(endpoints):
        row = endpoint_table.rows[i + 1].cells
        row[0].text = path
        row[1].text = method
        row[2].text = desc
        row[3].text = priority

    doc.add_heading("2.2 Field Mappings", level=2)
    doc.add_paragraph(
        "The following fields must be mapped from the LendFlow internal schema to the CIBIL "
        "request format:"
    )

    field_table = doc.add_table(rows=9, cols=4)
    field_table.style = "Light Grid Accent 1"
    fhdr = field_table.rows[0].cells
    fhdr[0].text = "LendFlow Field"
    fhdr[1].text = "CIBIL Field"
    fhdr[2].text = "Type"
    fhdr[3].text = "Required"

    fields = [
        ("applicant.pan_number", "pan_number", "string", "Yes"),
        ("applicant.full_name", "full_name", "string", "Yes"),
        ("applicant.date_of_birth", "date_of_birth", "date (YYYY-MM-DD)", "Yes"),
        ("applicant.mobile", "mobile_number", "string (10 digits)", "No"),
        ("applicant.email", "email_address", "string", "No"),
        ("applicant.address.full", "address", "string", "No"),
        ("loan.product_type", "loan_type", "string (enum)", "Yes"),
        ("loan.requested_amount", "loan_amount", "number", "Yes"),
    ]
    for i, (src, tgt, dtype, req) in enumerate(fields):
        row = field_table.rows[i + 1].cells
        row[0].text = src
        row[1].text = tgt
        row[2].text = dtype
        row[3].text = req

    doc.add_heading("2.3 Response Handling", level=2)
    doc.add_paragraph(
        "The CIBIL response must be parsed and mapped back to the LendFlow credit assessment "
        "model. Key response fields include: credit_score (integer 300-900), score_range "
        "(categorical), enquiry_id (for tracking), report_id (for detailed report reference), "
        "active_accounts (count), and overdue_accounts (count). Scores below 650 trigger "
        "additional verification workflows."
    )

    # -- 3. Authentication Requirements --
    doc.add_heading("3. Authentication Requirements", level=1)
    doc.add_paragraph(
        "CIBIL API v1 uses a combination of API key and client certificate (mTLS) for "
        "authentication. The API key is passed in the X-API-Key header, and a client "
        "certificate signed by CIBIL's CA must be presented during TLS handshake."
    )

    doc.add_heading("3.1 Credential Management", level=2)
    cred_items = [
        "API keys must be stored in AWS Secrets Manager with automatic rotation every 90 days",
        "Client certificates must be stored in a dedicated certificate vault",
        "Certificate expiry monitoring with 30-day advance alerting",
        "Separate credentials for UAT and Production environments",
        "IP whitelisting required for production access (provided by CIBIL)",
    ]
    for item in cred_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3.2 OAuth2 Migration (v2)", level=2)
    doc.add_paragraph(
        "CIBIL API v2 migrates to OAuth2 client credentials flow. The integration must "
        "support both authentication methods during the transition period (estimated 6 months). "
        "Token refresh must be handled automatically with a 5-minute buffer before expiry."
    )

    # -- 4. Security Requirements --
    doc.add_heading("4. Security Requirements", level=1)

    security_items = [
        "All PII data (PAN, Aadhaar, mobile) must be encrypted at rest using AES-256",
        "Data in transit must use TLS 1.2 or higher",
        "PAN numbers must be masked in all logs (show only last 4 characters)",
        "Credit reports must not be cached for more than 24 hours",
        "Access to credit bureau APIs must be logged with full audit trail",
        "Consent records must be maintained for minimum 5 years per RBI guidelines",
        "Rate limiting: maximum 100 requests per second per tenant",
        "DDoS protection with circuit breaker pattern (50% failure = circuit open)",
    ]
    for item in security_items:
        doc.add_paragraph(item, style="List Bullet")

    # -- 5. SLA Requirements --
    doc.add_heading("5. SLA Requirements", level=1)

    sla_table = doc.add_table(rows=7, cols=3)
    sla_table.style = "Light Grid Accent 1"
    sla_hdr = sla_table.rows[0].cells
    sla_hdr[0].text = "Metric"
    sla_hdr[1].text = "Target"
    sla_hdr[2].text = "Measurement"

    slas = [
        ("API Response Time (p50)", "< 500ms", "Per-request latency"),
        ("API Response Time (p99)", "< 2000ms", "Per-request latency"),
        ("Availability", "99.95%", "Monthly uptime"),
        ("Error Rate", "< 0.1%", "Non-timeout errors"),
        ("Throughput", "500 req/min", "Sustained load"),
        ("Recovery Time (RTO)", "< 5 minutes", "From circuit breaker open to close"),
    ]
    for i, (metric, target, measurement) in enumerate(slas):
        row = sla_table.rows[i + 1].cells
        row[0].text = metric
        row[1].text = target
        row[2].text = measurement

    doc.add_heading("5.1 Retry Policy", level=2)
    doc.add_paragraph(
        "Transient failures must be retried with exponential backoff: initial delay 100ms, "
        "max 3 retries, backoff multiplier 2x, max delay 2 seconds. Non-retryable errors "
        "(4xx responses) must fail immediately and be logged."
    )

    doc.add_heading("5.2 Fallback Strategy", level=2)
    doc.add_paragraph(
        "If CIBIL is unreachable after retries, the system must: (1) Return a cached score "
        "if available and less than 24 hours old, (2) Queue the request for async processing, "
        "(3) Notify the operations team via PagerDuty alert, (4) Log the incident for SLA "
        "tracking."
    )

    # -- 6. Data Transformation Rules --
    doc.add_heading("6. Data Transformation Rules", level=1)
    doc.add_paragraph(
        "The following transformations must be applied during request/response processing:"
    )

    transform_items = [
        "PAN number: Validate format (AAAAA9999A) before sending",
        "Date of birth: Convert from LendFlow format (DD/MM/YYYY) to CIBIL format (YYYY-MM-DD)",
        "Mobile number: Strip country code prefix (+91) if present",
        "Loan amount: Convert from paisa (internal) to rupees (CIBIL expects whole numbers)",
        "Address: Concatenate address_line1, address_line2, city, state, pincode",
        "Credit score: Normalize to 0-100 scale for internal scoring model",
    ]
    for item in transform_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(str(output_path))
    return output_path


# ──────────────────────────────────────────────────────────────────────────────
# API Client
# ──────────────────────────────────────────────────────────────────────────────

class FinSparkClient:
    """Thin wrapper around the FinSpark API."""

    def __init__(self, base_url: str, tenant_id: str = "demo-tenant", tenant_name: str = "Demo Bank") -> None:
        self.base_url = base_url.rstrip("/")
        self.api = f"{self.base_url}/api/v1"
        self.headers = {
            "X-Tenant-ID": tenant_id,
            "X-Tenant-Name": tenant_name,
            "X-Tenant-Role": "admin",
        }
        self.client = httpx.Client(headers=self.headers, timeout=30.0)

    def health(self) -> dict[str, Any]:
        resp = self.client.get(f"{self.base_url}/health")
        resp.raise_for_status()
        return resp.json()

    def upload_document(self, file_path: Path, doc_type: str = "brd") -> dict[str, Any]:
        with open(file_path, "rb") as f:
            resp = self.client.post(
                f"{self.api}/documents/upload",
                params={"doc_type": doc_type},
                files={"file": (file_path.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
        resp.raise_for_status()
        return resp.json()

    def get_document(self, doc_id: str) -> dict[str, Any]:
        resp = self.client.get(f"{self.api}/documents/{doc_id}")
        resp.raise_for_status()
        return resp.json()

    def list_adapters(self, category: str | None = None) -> dict[str, Any]:
        params = {}
        if category:
            params["category"] = category
        resp = self.client.get(f"{self.api}/adapters/", params=params)
        resp.raise_for_status()
        return resp.json()

    def get_adapter(self, adapter_id: str) -> dict[str, Any]:
        resp = self.client.get(f"{self.api}/adapters/{adapter_id}")
        resp.raise_for_status()
        return resp.json()

    def generate_config(self, document_id: str, adapter_version_id: str, name: str) -> dict[str, Any]:
        resp = self.client.post(
            f"{self.api}/configurations/generate",
            json={
                "document_id": document_id,
                "adapter_version_id": adapter_version_id,
                "name": name,
                "auto_map": True,
            },
        )
        resp.raise_for_status()
        return resp.json()

    def validate_config(self, config_id: str) -> dict[str, Any]:
        resp = self.client.post(f"{self.api}/configurations/{config_id}/validate")
        resp.raise_for_status()
        return resp.json()

    def run_simulation(self, configuration_id: str, test_type: str = "full") -> dict[str, Any]:
        resp = self.client.post(
            f"{self.api}/simulations/run",
            json={"configuration_id": configuration_id, "test_type": test_type},
        )
        resp.raise_for_status()
        return resp.json()

    def get_audit_logs(self, page: int = 1, page_size: int = 20) -> dict[str, Any]:
        resp = self.client.get(
            f"{self.api}/audit/",
            params={"page": page, "page_size": page_size},
        )
        resp.raise_for_status()
        return resp.json()

    def close(self) -> None:
        self.client.close()


# ──────────────────────────────────────────────────────────────────────────────
# Demo Flow
# ──────────────────────────────────────────────────────────────────────────────

def run_demo(base_url: str) -> None:
    banner()
    client = FinSparkClient(base_url)

    try:
        _run_all_steps(client)
    except httpx.ConnectError:
        error(f"Cannot connect to {base_url}")
        error("Make sure the backend is running: make dev  OR  uvicorn finspark.main:app --reload")
        sys.exit(1)
    except httpx.HTTPStatusError as e:
        error(f"HTTP {e.response.status_code}: {e.response.text[:500]}")
        sys.exit(1)
    except KeyboardInterrupt:
        print(f"\n{YELLOW}Demo interrupted.{RESET}")
        sys.exit(0)
    finally:
        client.close()


def _run_all_steps(client: FinSparkClient) -> None:
    # ── Step 0: Health Check ──────────────────────────────────────────────
    step_header(0, "Health Check")
    t0 = time.monotonic()
    health = client.health()
    elapsed = (time.monotonic() - t0) * 1000
    success(f"Server is {GREEN}{health['status']}{RESET} (v{health['version']}) [{elapsed:.0f}ms]")
    info(f"AI enabled: {health.get('checks', {}).get('ai_enabled', 'unknown')}")

    # ── Step 1: Create & Upload BRD ──────────────────────────────────────
    step_header(1, "Create & Upload CIBIL Bureau BRD")

    with tempfile.TemporaryDirectory() as tmpdir:
        brd_path = Path(tmpdir) / "CIBIL_Bureau_Integration_BRD.docx"
        info("Generating BRD document with python-docx...")
        create_brd_docx(brd_path)
        file_size_kb = brd_path.stat().st_size / 1024
        success(f"Created BRD: {brd_path.name} ({file_size_kb:.1f} KB)")

        info("Uploading to FinSpark API...")
        t0 = time.monotonic()
        upload_resp = client.upload_document(brd_path, doc_type="brd")
        elapsed = (time.monotonic() - t0) * 1000

    upload_data = upload_resp["data"]
    doc_id = upload_data["id"]
    status = upload_data["status"]

    if status == "parsed":
        success(f"Document parsed successfully [{elapsed:.0f}ms]")
    else:
        warn(f"Document status: {status}")

    print_table(
        ["Field", "Value"],
        [
            ["Document ID", doc_id],
            ["Filename", upload_data["filename"]],
            ["Type", upload_data["doc_type"]],
            ["Status", status],
        ],
    )

    # ── Step 2: View Parsed Results ──────────────────────────────────────
    step_header(2, "View Parsed Document Results")

    doc_detail = client.get_document(doc_id)
    parsed = doc_detail["data"].get("parsed_result")

    if parsed:
        success(f"Title: {parsed.get('title', 'N/A')}")
        info(f"Confidence: {parsed.get('confidence_score', 0):.0%}")

        services = parsed.get("services_identified", [])
        if services:
            info(f"Services identified: {', '.join(services)}")

        endpoints = parsed.get("endpoints", [])
        if endpoints:
            print(f"\n  {BOLD}Extracted Endpoints:{RESET}")
            ep_rows = []
            for ep in endpoints:
                ep_rows.append([ep.get("method", "?"), ep.get("path", "?"), ep.get("description", "")[:40]])
            print_table(["Method", "Path", "Description"], ep_rows)

        fields = parsed.get("fields", [])
        if fields:
            print(f"\n  {BOLD}Extracted Fields ({len(fields)}):{RESET}")
            field_rows = []
            for f in fields[:10]:
                req = f"{GREEN}Yes{RESET}" if f.get("is_required") else f"{DIM}No{RESET}"
                field_rows.append([f.get("name", "?"), f.get("data_type", "?"), req])
            print_table(["Field Name", "Type", "Required"], field_rows)
            if len(fields) > 10:
                info(f"... and {len(fields) - 10} more fields")

        auth_reqs = parsed.get("auth_requirements", [])
        if auth_reqs:
            print(f"\n  {BOLD}Authentication:{RESET}")
            for auth in auth_reqs:
                info(f"Type: {auth.get('auth_type', 'unknown')}")

        security = parsed.get("security_requirements", [])
        if security:
            print(f"\n  {BOLD}Security Requirements ({len(security)}):{RESET}")
            for s in security[:5]:
                print(f"  {DIM}│{RESET} {YELLOW}●{RESET} {s[:80]}")

        sla = parsed.get("sla_requirements", {})
        if sla:
            print(f"\n  {BOLD}SLA Requirements:{RESET}")
            sla_rows = [[k, v] for k, v in sla.items()]
            if sla_rows:
                print_table(["Metric", "Target"], sla_rows[:8])
    else:
        warn("No parsed result available (AI parsing may be disabled)")
        info("Continuing with adapter selection...")

    # ── Step 3: List Available Adapters ───────────────────────────────────
    step_header(3, "Browse Available Integration Adapters")

    adapters_resp = client.list_adapters()
    adapters_data = adapters_resp["data"]
    adapters_list = adapters_data.get("adapters", [])
    categories = adapters_data.get("categories", [])

    success(f"Found {adapters_data['total']} adapters across {len(categories)} categories")
    info(f"Categories: {', '.join(categories)}")

    adapter_rows = []
    cibil_adapter_id = None
    cibil_version_id = None

    for adapter in adapters_list:
        versions = adapter.get("versions", [])
        ver_str = ", ".join(v["version"] for v in versions)
        active = f"{GREEN}Active{RESET}" if adapter["is_active"] else f"{RED}Inactive{RESET}"
        adapter_rows.append([
            adapter["name"][:30],
            adapter["category"],
            ver_str,
            active,
        ])

        # Pick the CIBIL adapter for our demo
        if "cibil" in adapter["name"].lower():
            cibil_adapter_id = adapter["id"]
            if versions:
                cibil_version_id = versions[0]["id"]

    print_table(["Adapter", "Category", "Versions", "Status"], adapter_rows)

    if not cibil_adapter_id or not cibil_version_id:
        error("CIBIL adapter not found in registry. Is the backend seeded?")
        sys.exit(1)

    print(f"\n  {MAGENTA}▶{RESET} Selected: {BOLD}CIBIL Credit Bureau{RESET} (ID: {DIM}{cibil_adapter_id[:12]}...{RESET})")

    # Show adapter detail
    adapter_detail = client.get_adapter(cibil_adapter_id)
    adapter_info = adapter_detail["data"]
    versions = adapter_info.get("versions", [])

    for v in versions:
        print(f"\n  {BOLD}Version: {v['version']}{RESET} ({v['status']})")
        info(f"Auth: {v['auth_type']} | Base URL: {v.get('base_url', 'N/A')}")
        ep_list = v.get("endpoints", [])
        if ep_list:
            for ep in ep_list:
                print(f"    {CYAN}{ep['method']:6s}{RESET} {ep['path']}")

    # ── Step 4: Generate Configuration ────────────────────────────────────
    step_header(4, "Generate Integration Configuration")

    config_name = "CIBIL-Bureau-LendFlow-Integration"
    info(f"Mapping parsed BRD fields to CIBIL adapter schema...")
    info(f"Document: {doc_id[:12]}... -> Adapter Version: {cibil_version_id[:12]}...")

    t0 = time.monotonic()
    config_resp = client.generate_config(
        document_id=doc_id,
        adapter_version_id=cibil_version_id,
        name=config_name,
    )
    elapsed = (time.monotonic() - t0) * 1000

    config_data = config_resp["data"]
    config_id = config_data["id"]
    success(f"Configuration generated [{elapsed:.0f}ms]")

    print_table(
        ["Field", "Value"],
        [
            ["Config ID", config_id],
            ["Name", config_data["name"]],
            ["Status", config_data["status"]],
            ["Version", str(config_data["version"])],
        ],
    )

    mappings = config_data.get("field_mappings", [])
    if mappings:
        print(f"\n  {BOLD}Field Mappings ({len(mappings)}):{RESET}")
        mapping_rows = []
        for m in mappings:
            conf = m.get("confidence", 0)
            if conf >= 0.8:
                conf_str = f"{GREEN}{conf:.0%}{RESET}"
            elif conf >= 0.5:
                conf_str = f"{YELLOW}{conf:.0%}{RESET}"
            else:
                conf_str = f"{RED}{conf:.0%}{RESET}"
            transform = m.get("transformation") or "—"
            mapping_rows.append([
                m.get("source_field", "?"),
                m.get("target_field", "?"),
                transform,
                conf_str,
            ])
        print_table(["Source", "Target", "Transform", "Confidence"], mapping_rows)

    # ── Step 5: Validate Configuration ────────────────────────────────────
    step_header(5, "Validate Configuration")

    t0 = time.monotonic()
    validation = client.validate_config(config_id)
    elapsed = (time.monotonic() - t0) * 1000

    val_data = validation["data"]
    is_valid = val_data["is_valid"]
    coverage = val_data.get("coverage_score", 0)

    if is_valid:
        success(f"Configuration is {GREEN}{BOLD}VALID{RESET} [{elapsed:.0f}ms]")
    else:
        warn(f"Configuration has validation issues [{elapsed:.0f}ms]")

    # Coverage bar
    bar_width = 40
    filled = int(coverage * bar_width)
    bar = f"{'█' * filled}{'░' * (bar_width - filled)}"
    color = GREEN if coverage >= 0.8 else YELLOW if coverage >= 0.5 else RED
    print(f"\n  Coverage: {color}{bar}{RESET} {coverage:.0%}")

    if val_data.get("errors"):
        print(f"\n  {RED}Errors:{RESET}")
        for err in val_data["errors"]:
            error(err)

    if val_data.get("warnings"):
        print(f"\n  {YELLOW}Warnings:{RESET}")
        for w in val_data["warnings"]:
            warn(w)

    # ── Step 6: Run Simulation ────────────────────────────────────────────
    step_header(6, "Run Integration Simulation")

    info("Running full simulation suite against configuration...")
    t0 = time.monotonic()
    sim_resp = client.run_simulation(config_id, test_type="full")
    elapsed = (time.monotonic() - t0) * 1000

    sim_data = sim_resp["data"]
    sim_status = sim_data["status"]
    total_tests = sim_data["total_tests"]
    passed_tests = sim_data["passed_tests"]
    failed_tests = sim_data["failed_tests"]
    duration = sim_data.get("duration_ms", 0)

    status_color = GREEN if sim_status == "passed" else RED
    print(f"\n  {BG_GREEN if sim_status == 'passed' else BG_RED}{WHITE}{BOLD} {sim_status.upper()} {RESET}")

    print_table(
        ["Metric", "Value"],
        [
            ["Simulation ID", sim_data["id"]],
            ["Status", f"{status_color}{sim_status}{RESET}"],
            ["Total Tests", str(total_tests)],
            ["Passed", f"{GREEN}{passed_tests}{RESET}"],
            ["Failed", f"{RED}{failed_tests}{RESET}" if failed_tests > 0 else f"{GREEN}0{RESET}"],
            ["Duration", f"{duration}ms"],
            ["Test Type", sim_data.get("test_type", "full")],
        ],
    )

    steps = sim_data.get("steps", [])
    if steps:
        print(f"\n  {BOLD}Test Steps:{RESET}")
        for step in steps:
            step_status = step["status"]
            icon = f"{GREEN}✓{RESET}" if step_status == "passed" else f"{RED}✗{RESET}"
            conf = step.get("confidence_score", 0)
            print(f"  {icon} {step['step_name']:<40s} {DIM}{step.get('duration_ms', 0):>6d}ms{RESET}  conf={conf:.0%}")
            if step.get("error_message"):
                print(f"    {RED}└─ {step['error_message']}{RESET}")

    # Also run a smoke test
    print(f"\n  {DIM}Running smoke test...{RESET}")
    smoke_resp = client.run_simulation(config_id, test_type="smoke")
    smoke_data = smoke_resp["data"]
    smoke_icon = f"{GREEN}✓{RESET}" if smoke_data["status"] == "passed" else f"{RED}✗{RESET}"
    print(f"  {smoke_icon} Smoke test: {smoke_data['passed_tests']}/{smoke_data['total_tests']} passed")

    # ── Step 7: Audit Trail ───────────────────────────────────────────────
    step_header(7, "Review Audit Trail")

    audit_resp = client.get_audit_logs(page_size=10)
    audit_data = audit_resp["data"]
    audit_items = audit_data.get("items", [])
    total_logs = audit_data.get("total", 0)

    success(f"Total audit entries: {total_logs}")

    if audit_items:
        audit_rows = []
        for entry in audit_items:
            ts = entry.get("created_at", "?")
            if isinstance(ts, str) and len(ts) > 19:
                ts = ts[:19]
            audit_rows.append([
                ts,
                entry.get("action", "?"),
                entry.get("resource_type", "?"),
                entry.get("actor", "?"),
            ])
        print_table(["Timestamp", "Action", "Resource", "Actor"], audit_rows)

    # ── Summary ───────────────────────────────────────────────────────────
    print(f"\n{'═' * 70}")
    print(f"{BG_BLUE}{WHITE}{BOLD}  DEMO COMPLETE  {RESET}")
    print(f"{'═' * 70}")
    print(f"""
  {BOLD}Workflow Summary:{RESET}
    {GREEN}✓{RESET} BRD document created and uploaded
    {GREEN}✓{RESET} Document parsed — fields, endpoints, auth, SLA extracted
    {GREEN}✓{RESET} CIBIL Credit Bureau adapter selected from registry
    {GREEN}✓{RESET} Integration configuration auto-generated with field mappings
    {GREEN}✓{RESET} Configuration validated (coverage: {coverage:.0%})
    {GREEN}✓{RESET} Simulation executed: {passed_tests}/{total_tests} tests passed
    {GREEN}✓{RESET} Full audit trail captured

  {BOLD}Key IDs:{RESET}
    Document:      {DIM}{doc_id}{RESET}
    Configuration: {DIM}{config_id}{RESET}
    Simulation:    {DIM}{sim_data['id']}{RESET}

  {BOLD}API Docs:{RESET}  {CYAN}{client.base_url}/docs{RESET}
  {BOLD}ReDoc:{RESET}     {CYAN}{client.base_url}/redoc{RESET}
""")


# ──────────────────────────────────────────────────────────────────────────────
# Entry point
# ──────────────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(description="FinSpark Platform Demo")
    parser.add_argument("--base-url", default="http://localhost:8000", help="Backend base URL")
    args = parser.parse_args()
    run_demo(args.base_url)


if __name__ == "__main__":
    main()
