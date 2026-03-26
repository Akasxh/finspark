"""
DOCX parser using python-docx.

Extracts:
- Hierarchical sections (Heading 1/2/3 → DocumentSection tree)
- Tables (headers + rows)
- Inline URLs, API paths, field names per section
- Auth-scheme hints from content
- Rough endpoint mentions from inline text
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Union

from docx import Document  # type: ignore[import-untyped]
from docx.oxml.ns import qn  # type: ignore[import-untyped]
from docx.table import Table  # type: ignore[import-untyped]
from docx.text.paragraph import Paragraph  # type: ignore[import-untyped]

from finspark.models.parsed_document import (
    ApiEndpoint,
    AuthRequirement,
    AuthScheme,
    DocumentSection,
    DocumentType,
    FieldDefinition,
    HttpMethod,
    ParsedDocument,
    TableData,
)
from finspark.services.document_parser.heuristics import (
    classify_section,
    detect_auth_schemes,
    detect_http_method_pairs,
    extract_api_paths,
    extract_field_names,
    extract_urls,
    extract_version,
)


def _heading_level(para: Paragraph) -> int | None:
    """Return heading level (1-9) or None if not a heading paragraph."""
    style_name: str = para.style.name if para.style else ""
    if style_name.startswith("Heading"):
        try:
            return int(style_name.split()[-1])
        except ValueError:
            return 1
    # Fallback: check outline level in paragraph XML
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        outlineLvl = pPr.find(qn("w:outlineLvl"))
        if outlineLvl is not None:
            val = outlineLvl.get(qn("w:val"))
            if val is not None:
                return int(val) + 1
    return None


def _extract_table(tbl: Table, section_heading: str) -> TableData:
    rows_data: list[list[str]] = []
    for row in tbl.rows:
        rows_data.append([cell.text.strip() for cell in row.cells])

    headers: list[str] = rows_data[0] if rows_data else []
    body_rows = rows_data[1:] if len(rows_data) > 1 else []
    return TableData(
        headers=headers,
        rows=body_rows,
        source_section=section_heading,
    )


def _section_from_stack(
    heading: str,
    level: int,
    paragraphs: list[str],
    tables: list[TableData],
) -> DocumentSection:
    content = "\n".join(paragraphs)
    return DocumentSection(
        heading=heading,
        level=level,
        category=classify_section(heading, content),
        content=content,
        tables=tables,
        urls=extract_urls(content),
        api_paths=extract_api_paths(content),
        field_names=extract_field_names(content),
    )


def parse_docx(source: Union[str, Path, bytes, io.BytesIO]) -> ParsedDocument:
    if isinstance(source, bytes):
        source = io.BytesIO(source)

    path_str = str(source) if not isinstance(source, io.BytesIO) else "<bytes>"

    try:
        doc = Document(source)
    except Exception as exc:
        return ParsedDocument(
            source_filename=path_str,
            doc_type=DocumentType.DOCX,
            parse_errors=[f"Failed to open DOCX: {exc}"],
        )

    # -----------------------------------------------------------------------
    # Walk document body — paragraphs and inline tables share an element order
    # -----------------------------------------------------------------------
    # We maintain a stack of (level, heading, [para_texts], [tables])
    SectionFrame = tuple[int, str, list[str], list[TableData]]

    stack: list[SectionFrame] = []
    # Current accumulator for the section under construction
    cur_level: int = 0
    cur_heading: str = "Preamble"
    cur_paras: list[str] = []
    cur_tables: list[TableData] = []
    flat_sections: list[DocumentSection] = []
    all_raw_text_parts: list[str] = []

    def flush_current() -> None:
        if cur_paras or cur_tables:
            flat_sections.append(
                _section_from_stack(cur_heading, cur_level, cur_paras, cur_tables)
            )

    # Iterate body XML children to preserve paragraph/table order
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            para = Paragraph(child, doc)
            text = para.text.strip()
            if not text:
                continue
            all_raw_text_parts.append(text)

            lvl = _heading_level(para)
            if lvl is not None:
                flush_current()
                cur_level = lvl
                cur_heading = text
                cur_paras = []
                cur_tables = []
            else:
                cur_paras.append(text)

        elif tag == "tbl":
            tbl = Table(child, doc)
            extracted = _extract_table(tbl, cur_heading)
            cur_tables.append(extracted)
            # Also flatten table text for raw_text
            for row in [extracted.headers] + extracted.rows:
                all_raw_text_parts.append(" | ".join(row))

    flush_current()

    # -----------------------------------------------------------------------
    # Build top-level structure  (flat_sections → nested tree not needed for
    # hackathon; flat list keyed by level is enough for consumers)
    # -----------------------------------------------------------------------
    raw_text = "\n".join(all_raw_text_parts)
    title = flat_sections[0].heading if flat_sections else ""

    # Global entity extraction across full text
    all_urls = extract_urls(raw_text)
    all_api_paths = extract_api_paths(raw_text)
    all_field_names = extract_field_names(raw_text)
    version = extract_version(raw_text)

    # Auth requirements
    auth_schemes = detect_auth_schemes(raw_text)
    auth_requirements = [
        AuthRequirement(
            scheme=scheme,
            description=f"Detected via heuristic scan of document text",
        )
        for scheme in auth_schemes
    ] if auth_schemes else [AuthRequirement(scheme=AuthScheme.NONE)]

    # Endpoint mentions from free text
    endpoints: list[ApiEndpoint] = []
    seen_endpoints: set[tuple[str, str]] = set()
    for section in flat_sections:
        for method_str, path in detect_http_method_pairs(section.content):
            key = (method_str, path)
            if key not in seen_endpoints:
                seen_endpoints.add(key)
                try:
                    method = HttpMethod(method_str)
                except ValueError:
                    continue
                endpoints.append(
                    ApiEndpoint(
                        path=path,
                        method=method,
                        source_section=section.heading,
                        auth_required=bool(auth_schemes),
                    )
                )

    # Field definitions from tables with "field" / "name" / "type" headers
    field_definitions: list[FieldDefinition] = []
    all_tables: list[TableData] = []
    for section in flat_sections:
        for tbl in section.tables:
            all_tables.append(tbl)
            lower_headers = [h.lower() for h in tbl.headers]
            name_idx = next(
                (i for i, h in enumerate(lower_headers) if h in ("field", "name", "property", "parameter")),
                None,
            )
            type_idx = next(
                (i for i, h in enumerate(lower_headers) if "type" in h),
                None,
            )
            desc_idx = next(
                (i for i, h in enumerate(lower_headers) if "desc" in h or "note" in h),
                None,
            )
            req_idx = next(
                (i for i, h in enumerate(lower_headers) if "req" in h or "mandatory" in h),
                None,
            )
            if name_idx is None:
                continue
            for row in tbl.rows:
                if len(row) <= name_idx:
                    continue
                fdef = FieldDefinition(
                    name=row[name_idx],
                    field_type=row[type_idx] if type_idx is not None and len(row) > type_idx else "unknown",
                    description=row[desc_idx] if desc_idx is not None and len(row) > desc_idx else "",
                    required=bool(
                        req_idx is not None
                        and len(row) > req_idx
                        and row[req_idx].lower() in ("yes", "true", "required", "mandatory", "y")
                    ),
                )
                field_definitions.append(fdef)

    return ParsedDocument(
        source_filename=path_str,
        doc_type=DocumentType.DOCX,
        title=title,
        version=version,
        sections=flat_sections,
        tables=all_tables,
        endpoints=endpoints,
        auth_requirements=auth_requirements,
        field_definitions=field_definitions,
        all_urls=all_urls,
        all_api_paths=all_api_paths,
        all_field_names=all_field_names,
        word_count=len(raw_text.split()),
        raw_text=raw_text,
    )
