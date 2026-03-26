"""
Unit tests for finspark.services.document_parser and finspark.models.parsed_document.

Covers heuristics, DOCX parser, PDF parser, OpenAPI parser, and the facade.
No filesystem I/O — all sources are in-memory bytes.
"""
from __future__ import annotations

import io
import json
from typing import Any

import pytest

from finspark.models.parsed_document import (
    AuthScheme,
    DocumentType,
    HttpMethod,
    SectionCategory,
)
from finspark.services.document_parser import parse_document_bytes
from finspark.services.document_parser.heuristics import (
    classify_section,
    detect_auth_schemes,
    detect_http_method_pairs,
    extract_api_paths,
    extract_field_names,
    extract_urls,
    extract_version,
)
from finspark.services.document_parser.openapi_parser import parse_openapi


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_docx_bytes(paragraphs: list[tuple[str, str | None]]) -> bytes:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    for text, style in paragraphs:
        if style:
            doc.add_heading(text, level=int(style.split()[-1]))
        else:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _openapi3_spec(extra_paths: dict[str, Any] | None = None) -> dict[str, Any]:
    spec: dict[str, Any] = {
        "openapi": "3.0.3",
        "info": {"title": "Test API", "version": "1.2.3"},
        "servers": [{"url": "https://api.example.com/v1"}],
        "paths": {
            "/users": {
                "get": {
                    "summary": "List users",
                    "tags": ["users"],
                    "security": [{"BearerAuth": []}],
                    "responses": {"200": {"description": "OK"}},
                },
                "post": {
                    "summary": "Create user",
                    "tags": ["users"],
                    "requestBody": {
                        "content": {
                            "application/json": {
                                "schema": {"$ref": "#/components/schemas/User"}
                            }
                        }
                    },
                    "responses": {"201": {"description": "Created"}},
                },
            },
            "/users/{id}": {
                "get": {
                    "summary": "Get user by id",
                    "tags": ["users"],
                    "parameters": [
                        {"name": "id", "in": "path", "required": True, "schema": {"type": "string"}}
                    ],
                    "responses": {"200": {"description": "OK"}},
                }
            },
        },
        "components": {
            "schemas": {
                "User": {
                    "type": "object",
                    "required": ["user_id", "email"],
                    "properties": {
                        "user_id": {"type": "string", "format": "uuid"},
                        "email": {"type": "string", "format": "email"},
                        "first_name": {"type": "string"},
                        "last_name": {"type": "string"},
                        "created_at": {"type": "string", "format": "date-time"},
                    },
                }
            },
            "securitySchemes": {
                "BearerAuth": {"type": "http", "scheme": "bearer"},
            },
        },
    }
    if extra_paths:
        spec["paths"].update(extra_paths)
    return spec


# ---------------------------------------------------------------------------
# Heuristics
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_extract_urls_basic() -> None:
    text = "See https://api.example.com/v1 and http://foo.bar/baz?x=1 for details."
    urls = extract_urls(text)
    assert "https://api.example.com/v1" in urls
    assert "http://foo.bar/baz?x=1" in urls


@pytest.mark.unit
def test_extract_api_paths() -> None:
    text = "Call POST /api/v2/payments and GET /v1/accounts/{accountId}/balance"
    paths = extract_api_paths(text)
    assert any("/api/v2/payments" in p for p in paths)
    assert any("/v1/accounts/" in p for p in paths)


@pytest.mark.unit
def test_extract_field_names_snake_and_camel() -> None:
    text = "Fields: account_id, transaction_type, referenceNumber, paymentMethod"
    names = extract_field_names(text)
    assert "account_id" in names
    assert "transaction_type" in names
    assert "referenceNumber" in names
    assert "paymentMethod" in names


@pytest.mark.unit
def test_extract_version() -> None:
    assert extract_version("API Version: 2.3.1") == "2.3.1"
    assert extract_version("v1.0") == "1.0"
    assert extract_version("no version here") == ""


@pytest.mark.unit
def test_classify_section_requirements() -> None:
    cat = classify_section("Functional Requirements", "The system must support OAuth2")
    assert cat == SectionCategory.REQUIREMENTS


@pytest.mark.unit
def test_classify_section_security() -> None:
    cat = classify_section("Security Considerations", "TLS encryption is mandatory")
    assert cat == SectionCategory.SECURITY


@pytest.mark.unit
def test_classify_section_authentication() -> None:
    cat = classify_section("Authentication", "Bearer token via JWT")
    assert cat == SectionCategory.AUTHENTICATION


@pytest.mark.unit
def test_detect_auth_schemes_bearer() -> None:
    schemes = detect_auth_schemes("Use Bearer token in Authorization header")
    assert AuthScheme.BEARER in schemes


@pytest.mark.unit
def test_detect_auth_schemes_multi() -> None:
    text = "Supports OAuth2 and API key via x-api-key header"
    schemes = detect_auth_schemes(text)
    assert AuthScheme.OAUTH2 in schemes
    assert AuthScheme.API_KEY in schemes


@pytest.mark.unit
def test_detect_http_method_pairs() -> None:
    text = "Use GET /api/users or POST /api/users to list or create"
    pairs = detect_http_method_pairs(text)
    methods = {p[0] for p in pairs}
    assert "GET" in methods
    assert "POST" in methods


# ---------------------------------------------------------------------------
# DOCX parser
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_docx_parse_basic_sections() -> None:
    content = _make_docx_bytes([
        ("API Integration Guide", "Heading 1"),
        ("This document describes the payment integration requirements.", None),
        ("Authentication", "Heading 2"),
        ("Use Bearer token in the Authorization header.", None),
        ("Endpoints", "Heading 2"),
        ("POST /api/v1/payments creates a new payment.", None),
    ])
    result = parse_document_bytes(content, "guide.docx")

    assert result.doc_type == DocumentType.DOCX
    assert len(result.sections) >= 3
    assert result.word_count > 0
    assert any(s.heading == "Authentication" for s in result.sections)


@pytest.mark.unit
def test_docx_extracts_urls() -> None:
    content = _make_docx_bytes([
        ("Overview", "Heading 1"),
        ("See https://docs.example.com/api for full reference.", None),
    ])
    result = parse_document_bytes(content, "spec.docx")
    assert "https://docs.example.com/api" in result.all_urls


@pytest.mark.unit
def test_docx_detects_bearer_auth() -> None:
    content = _make_docx_bytes([
        ("Security", "Heading 1"),
        ("All requests must include a Bearer token in the Authorization header.", None),
    ])
    result = parse_document_bytes(content, "brd.docx")
    schemes = {r.scheme for r in result.auth_requirements}
    assert AuthScheme.BEARER in schemes


@pytest.mark.unit
def test_docx_endpoint_extraction() -> None:
    content = _make_docx_bytes([
        ("API", "Heading 1"),
        ("POST /api/v1/transactions — create transaction", None),
        ("GET /api/v1/transactions/{id} — fetch single transaction", None),
    ])
    result = parse_document_bytes(content, "api.docx")
    paths = {ep.path for ep in result.endpoints}
    methods = {ep.method for ep in result.endpoints}
    assert "/api/v1/transactions" in paths
    assert HttpMethod.POST in methods
    assert HttpMethod.GET in methods


# ---------------------------------------------------------------------------
# OpenAPI parser
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_openapi_json_parse() -> None:
    spec = _openapi3_spec()
    content = json.dumps(spec).encode()
    result = parse_document_bytes(content, "openapi.json")

    assert result.doc_type == DocumentType.OPENAPI_JSON
    assert result.title == "Test API"
    assert result.version == "1.2.3"
    assert result.openapi_version == "3.0.3"
    assert "https://api.example.com/v1" in result.base_urls


@pytest.mark.unit
def test_openapi_endpoints_extracted() -> None:
    spec = _openapi3_spec()
    result = parse_openapi(json.dumps(spec).encode(), filename="api.json")

    paths = {ep.path for ep in result.endpoints}
    assert "/users" in paths
    assert "/users/{id}" in paths

    methods = {ep.method for ep in result.endpoints}
    assert HttpMethod.GET in methods
    assert HttpMethod.POST in methods


@pytest.mark.unit
def test_openapi_auth_required_flag() -> None:
    spec = _openapi3_spec()
    result = parse_openapi(json.dumps(spec).encode(), filename="api.json")

    get_users = next(
        (ep for ep in result.endpoints if ep.path == "/users" and ep.method == HttpMethod.GET),
        None,
    )
    assert get_users is not None
    assert get_users.auth_required is True


@pytest.mark.unit
def test_openapi_field_definitions() -> None:
    spec = _openapi3_spec()
    result = parse_openapi(json.dumps(spec).encode(), filename="api.json")

    field_names = {fd.name for fd in result.field_definitions}
    assert "User.user_id" in field_names
    assert "User.email" in field_names
    assert "User.first_name" in field_names

    user_id = next(fd for fd in result.field_definitions if fd.name == "User.user_id")
    assert user_id.required is True


@pytest.mark.unit
def test_openapi_security_scheme() -> None:
    spec = _openapi3_spec()
    result = parse_openapi(json.dumps(spec).encode(), filename="api.json")

    schemes = {r.scheme for r in result.auth_requirements}
    assert AuthScheme.BEARER in schemes


@pytest.mark.unit
def test_openapi_yaml_parse() -> None:
    import yaml

    spec = _openapi3_spec()
    content = yaml.dump(spec).encode()
    result = parse_document_bytes(content, "openapi.yaml")

    assert result.doc_type == DocumentType.OPENAPI_YAML
    assert result.title == "Test API"
    assert len(result.endpoints) == 3


@pytest.mark.unit
def test_openapi_swagger2() -> None:
    spec: dict[str, Any] = {
        "swagger": "2.0",
        "info": {"title": "Swagger Test", "version": "0.1"},
        "host": "api.example.com",
        "basePath": "/v1",
        "schemes": ["https"],
        "paths": {
            "/payments": {
                "post": {
                    "summary": "Create payment",
                    "parameters": [],
                    "responses": {"200": {"description": "OK"}},
                }
            }
        },
        "securityDefinitions": {
            "ApiKeyAuth": {"type": "apiKey", "name": "X-API-Key", "in": "header"}
        },
    }
    content = json.dumps(spec).encode()
    result = parse_openapi(content, filename="swagger.json")

    assert result.title == "Swagger Test"
    assert "https://api.example.com/v1" in result.base_urls
    paths = {ep.path for ep in result.endpoints}
    assert "/payments" in paths
    schemes = {r.scheme for r in result.auth_requirements}
    assert AuthScheme.API_KEY in schemes


@pytest.mark.unit
def test_openapi_invalid_spec_returns_error() -> None:
    result = parse_openapi(b"this is not valid yaml or json: ::::", filename="bad.yaml")
    assert result.parse_errors


@pytest.mark.unit
def test_parse_document_bytes_unknown_type() -> None:
    result = parse_document_bytes(b"random garbage data", "unknown.bin")
    assert result is not None


@pytest.mark.unit
def test_parsed_document_model_copy_preserves_data() -> None:
    spec = _openapi3_spec()
    result = parse_openapi(json.dumps(spec).encode(), filename="api.json")
    copy = result.model_copy(update={"source_filename": "patched.json"})
    assert copy.source_filename == "patched.json"
    assert copy.title == result.title
    assert len(copy.endpoints) == len(result.endpoints)
